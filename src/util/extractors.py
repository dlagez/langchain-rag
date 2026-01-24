from __future__ import annotations

import os
from pathlib import Path
import shutil
import subprocess
import tempfile

from .file_utils import _read_text_with_fallback
from .text_utils import _ensure_page_markers, _maybe_repair_mojibake, _normalize_text


def _extract_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []
    page_number = 1
    parts.append(f"=== Page {page_number} ===")

    def add_text(text: str) -> None:
        text = text.strip()
        if text:
            parts.append(text)

    def has_page_break(paragraph) -> bool:
        element = getattr(paragraph, "_p", None)
        if element is None:
            return False
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        xpath = ".//w:br[@w:type='page'] | .//w:lastRenderedPageBreak"
        try:
            return bool(element.xpath(xpath, namespaces=ns))
        except TypeError:
            return bool(
                element.xpath(
                    ".//*[local-name()='br' and @*[local-name()='type']='page']"
                    " | .//*[local-name()='lastRenderedPageBreak']"
                )
            )

    for paragraph in doc.paragraphs:
        add_text(paragraph.text)
        if has_page_break(paragraph):
            page_number += 1
            parts.append(f"=== Page {page_number} ===")
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    text = _normalize_text("\n".join(parts))
    return _ensure_page_markers(_maybe_repair_mojibake(text))


def _extract_excel_rows(path: Path) -> list[dict]:
    """从 Excel 提取行级结构，包含 sheet、row、cells、text、is_header。"""
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None, dtype=str, header=None)
    rows: list[dict] = []
    for name, df in sheets.items():
        df = df.fillna("")
        header_row = None
        for idx, row in df.iterrows():
            cells = [str(val).strip() if val is not None else "" for val in row.tolist()]
            if all(cell == "" for cell in cells):
                continue
            text = "\t".join(cells).rstrip()
            if header_row is None:
                header_row = idx
                is_header = True
            else:
                is_header = False
            rows.append(
                {
                    "sheet": str(name),
                    "row": int(idx) + 1,
                    "cells": cells,
                    "text": text,
                    "is_header": is_header,
                }
            )
    return rows


def _excel_rows_to_text(rows: list[dict]) -> str:
    """将行级结构还原为可读文本，保留 Sheet 分隔。"""
    if not rows:
        return ""
    parts: list[str] = []
    current_sheet = None
    for row in rows:
        sheet = row.get("sheet")
        if sheet != current_sheet:
            current_sheet = sheet
            parts.append(f"[Sheet: {sheet}]")
        text = row.get("text") or ""
        if text:
            parts.append(text)
    return _normalize_text("\n".join(parts))


def _extract_excel_text(path: Path) -> str:
    rows = _extract_excel_rows(path)
    return _excel_rows_to_text(rows)


def _extract_doc_with_textract(path: Path) -> str | None:
    try:
        import textract  # type: ignore
    except ImportError:
        return None
    raw = textract.process(str(path))
    if not raw:
        return ""
    return raw.decode("utf-8", errors="replace")


def _extract_doc_with_word(path: Path) -> str | None:
    if os.name != "nt":
        return None
    try:
        import win32com.client as win32  # type: ignore
    except ImportError:
        return None
    word = None
    doc = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(path), ReadOnly=True)
        return doc.Content.Text
    except Exception:
        return None
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def _extract_doc_with_soffice(path: Path) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            tmpdir,
            str(path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            return None
        txt_path = Path(tmpdir) / f"{path.stem}.txt"
        if not txt_path.exists():
            candidates = list(Path(tmpdir).glob("*.txt"))
            if not candidates:
                return None
            txt_path = candidates[0]
        return _read_text_with_fallback(txt_path)


def _extract_doc_text(path: Path) -> str:
    for extractor in (
        _extract_doc_with_textract,
        _extract_doc_with_word,
        _extract_doc_with_soffice,
    ):
        text = extractor(path)
        if text is not None:
            return _ensure_page_markers(_maybe_repair_mojibake(_normalize_text(text)))
    raise RuntimeError(
        "Missing .doc extractor. Install textract/pywin32 or LibreOffice for .doc files."
    )
