from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

import numpy as np
from langchain_core.documents import Document
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, PointStruct, VectorParams

from ppocr_pdf_tool import LocalPPOCRTool, RemotePPOCRTool

_CJK_RE = re.compile(r"[\u4e00-\u9fff]+")
_WORD_RE = re.compile(r"[A-Za-z0-9]+")
_ATTACHMENT_RE = re.compile(r"附件\s*([0-9]+)")
_ATTACHMENT_CN_RE = re.compile(r"附件\s*([一二三四五六七八九十])")
_CN_NUM_MAP = {
    "一": "1",
    "二": "2",
    "三": "3",
    "四": "4",
    "五": "5",
    "六": "6",
    "七": "7",
    "八": "8",
    "九": "9",
    "十": "10",
}

logger = logging.getLogger(__name__)


def _cjk_ngrams(text: str, min_size: int = 2, max_size: int = 4) -> list[str]:
    ngrams: list[str] = []
    for chunk in _CJK_RE.findall(text):
        limit = min(max_size, len(chunk))
        for size in range(min_size, limit + 1):
            for idx in range(len(chunk) - size + 1):
                ngrams.append(chunk[idx : idx + size])
    seen = set()
    deduped: list[str] = []
    for term in ngrams:
        if term not in seen:
            seen.add(term)
            deduped.append(term)
    return deduped


def _latin_keywords(text: str) -> list[str]:
    tokens = [token.lower() for token in _WORD_RE.findall(text) if len(token) > 1]
    return list(dict.fromkeys(tokens))


def _query_terms(query: str) -> tuple[list[str], list[str]]:
    cjk = _cjk_ngrams(query, min_size=3, max_size=5)
    if not cjk:
        cjk = _cjk_ngrams(query, min_size=2, max_size=4)
    if len(cjk) > 64:
        cjk = cjk[:64]
    latin = _latin_keywords(query)
    if len(latin) > 32:
        latin = latin[:32]
    return cjk, latin


def _extract_source_hint(query: str) -> str | None:
    match = _ATTACHMENT_RE.search(query)
    if match:
        return f"附件{match.group(1)}"
    match = _ATTACHMENT_CN_RE.search(query)
    if match:
        number = _CN_NUM_MAP.get(match.group(1))
        if number:
            return f"附件{number}"
    return None


def _filter_docs_by_source_hint(
    docs: list[Document], source_hint: str | None
) -> list[Document]:
    if not source_hint:
        return docs
    filtered: list[Document] = []
    for doc in docs:
        source = doc.metadata.get("source") or ""
        if source_hint in Path(source).name:
            filtered.append(doc)
    return filtered or docs


def _keyword_score(text: str, cjk_keywords: list[str], latin_keywords: list[str]) -> int:
    score = 0
    for kw in cjk_keywords:
        if kw in text:
            score += len(kw)
    if latin_keywords:
        lowered = text.lower()
        for kw in latin_keywords:
            if kw in lowered:
                score += len(kw)
    return score


def _docs_have_keyword_hits(docs: list[Document], query: str) -> bool:
    cjk, latin = _query_terms(query)
    if not cjk and not latin:
        return True
    for doc in docs:
        if _keyword_score(doc.page_content, cjk, latin):
            return True
    return False


def _keyword_fallback(docs: list[Document], query: str, limit: int = 3) -> list[Document]:
    cjk, latin = _query_terms(query)
    if not cjk and not latin:
        return []
    scored = []
    for doc in docs:
        score = _keyword_score(doc.page_content, cjk, latin)
        if score:
            scored.append((score, doc))
    scored.sort(key=lambda item: (-item[0], item[1].metadata.get("page") or 0))
    return [doc for _, doc in scored[:limit]]


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _join_pages(pages: list[str]) -> str:
    chunks = []
    for idx, page_text in enumerate(pages, start=1):
        chunks.append(f"=== Page {idx} ===")
        chunks.append(page_text.strip())
    return "\n\n".join(chunks).strip()


def _split_pages(text: str, use_page_markers: bool = False) -> list[str]:
    if not use_page_markers or "=== Page " not in text:
        return [text.strip()] if text.strip() else []
    parts = []
    for block in text.split("=== Page "):
        block = block.strip()
        if not block:
            continue
        _, _, content = block.partition("===")
        parts.append(content.strip())
    return parts


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _read_text_with_fallback(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="gb18030")
        except UnicodeDecodeError:
            return path.read_text(encoding="utf-8", errors="replace")


def _build_metadata(path: Path, page: int | None = None) -> dict:
    metadata = {
        "source": str(path),
        "source_type": path.suffix.lower().lstrip("."),
    }
    if page is not None:
        metadata["page"] = page
    return metadata


def _docs_from_text(
    path: Path,
    text: str,
    use_page_markers: bool = False,
    force_page: bool = False,
) -> list[Document]:
    text = _normalize_text(text)
    if not text:
        return []
    pages = _split_pages(text, use_page_markers=use_page_markers)
    docs: list[Document] = []
    use_page = force_page or (use_page_markers and len(pages) > 1)
    for idx, page_text in enumerate(pages, start=1):
        page_text = page_text.strip()
        if not page_text:
            continue
        metadata = _build_metadata(path, page=idx if use_page else None)
        docs.append(Document(page_content=page_text, metadata=metadata))
    return docs


def _extract_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return _normalize_text("\n".join(parts))


def _extract_excel_text(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts = []
    for name, df in sheets.items():
        df = df.fillna("")
        parts.append(f"[Sheet: {name}]")
        parts.append(df.to_csv(sep="\t", index=False))
    return _normalize_text("\n".join(parts))


def _extract_doc_with_textract(path: Path) -> str | None:
    try:
        import textract  # type: ignore
    except ImportError:
        return None
    raw = textract.process(str(path))
    if not raw:
        return ""
    return raw.decode("utf-8", errors="replace")


def _extract_doc_with_word(path: Path) -> str | None:
    if os.name != "nt":
        return None
    try:
        import win32com.client as win32  # type: ignore
    except ImportError:
        return None
    word = None
    doc = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(path), ReadOnly=True)
        return doc.Content.Text
    except Exception:
        return None
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def _extract_doc_with_soffice(path: Path) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            tmpdir,
            str(path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            return None
        txt_path = Path(tmpdir) / f"{path.stem}.txt"
        if not txt_path.exists():
            candidates = list(Path(tmpdir).glob("*.txt"))
            if not candidates:
                return None
            txt_path = candidates[0]
        return _read_text_with_fallback(txt_path)


def _extract_doc_text(path: Path) -> str:
    for extractor in (
        _extract_doc_with_textract,
        _extract_doc_with_word,
        _extract_doc_with_soffice,
    ):
        text = extractor(path)
        if text is not None:
            return _normalize_text(text)
    raise RuntimeError(
        "Missing .doc extractor. Install textract/pywin32 or LibreOffice for .doc files."
    )


class _LazyOCR:
    def __init__(self, **kwargs) -> None:
        self._kwargs = kwargs
        self._ocr: LocalPPOCRTool | RemotePPOCRTool | None = None

    def get(self) -> LocalPPOCRTool | RemotePPOCRTool:
        if self._ocr is None:
            mode = os.getenv("PPOCR_MODE", "local").strip().lower()
            if mode in ("remote", "http"):
                self._ocr = RemotePPOCRTool()
            elif mode == "auto":
                try:
                    self._ocr = LocalPPOCRTool(**self._kwargs)
                except SystemExit:
                    self._ocr = RemotePPOCRTool()
            else:
                self._ocr = LocalPPOCRTool(**self._kwargs)
        return self._ocr


def _extract_text_from_file(
    path: Path,
    ocr: _LazyOCR,
    image_dpi: int = 200,
) -> tuple[list[Document], str]:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = _normalize_text(_read_text_with_fallback(path))
        return _docs_from_text(path, text), text
    if suffix == ".pdf":
        pages = LocalPPOCRTool._extract_pdf_text(path)
        if pages is None:
            start = time.perf_counter()
            pages = ocr.get().ocr_pdf(path, image_dpi=image_dpi)
            duration = time.perf_counter() - start
            logger.info("OCR processed %s in %.2fs", path, duration)
        text = _join_pages(pages)
        return _docs_from_text(path, text, use_page_markers=True, force_page=True), text
    if suffix == ".docx":
        text = _extract_docx_text(path)
        return _docs_from_text(path, text), text
    if suffix in (".xls", ".xlsx"):
        text = _extract_excel_text(path)
        return _docs_from_text(path, text), text
    if suffix == ".csv":
        text = _normalize_text(_read_text_with_fallback(path))
        return _docs_from_text(path, text), text
    if suffix in (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"):
        start = time.perf_counter()
        lines = ocr.get().ocr_image_path(path)
        duration = time.perf_counter() - start
        logger.info("OCR processed %s in %.2fs", path, duration)
        text = _normalize_text("\n".join(lines))
        return _docs_from_text(path, text), text
    if suffix == ".doc":
        text = _extract_doc_text(path)
        return _docs_from_text(path, text), text
    raise ValueError(f"Unsupported file type: {path}")


def _load_cached_text(source_path: Path, processed_path: Path) -> str | None:
    if not processed_path.exists():
        return None
    try:
        if processed_path.stat().st_mtime_ns >= source_path.stat().st_mtime_ns:
            return _normalize_text(
                processed_path.read_text(encoding="utf-8", errors="replace")
            )
    except OSError:
        return None
    return None


def process_sources(
    source_dir: Path,
    processed_dir: Path,
    ocr: _LazyOCR,
    image_dpi: int = 200,
) -> list[Document]:
    source_dir = Path(source_dir)
    processed_dir = Path(processed_dir)
    if not source_dir.exists():
        raise SystemExit(f"Source directory not found: {source_dir}")

    docs: list[Document] = []
    for path in sorted(source_dir.rglob("*")):
        if not path.is_file():
            continue

        rel = path.relative_to(source_dir)
        out_path = (processed_dir / rel).with_suffix(".txt")

        cached_text = _load_cached_text(path, out_path)
        if cached_text:
            cached_docs = _docs_from_text(
                path,
                cached_text,
                use_page_markers=path.suffix.lower() == ".pdf",
                force_page=path.suffix.lower() == ".pdf",
            )
            docs.extend(cached_docs)
            continue

        try:
            extracted_docs, text = _extract_text_from_file(
                path, ocr, image_dpi=image_dpi
            )
        except ValueError:
            logger.info("Skipping unsupported file: %s", path)
            continue
        except Exception as exc:
            logger.warning("Skipping %s: %s", path, exc)
            continue

        if text:
            _write_text(out_path, text)
        if extracted_docs:
            docs.extend(extracted_docs)

    if not docs:
        raise SystemExit("No supported files found in data/source.")
    return docs


def _fingerprint_processed(processed_dir: Path) -> str:
    if not processed_dir.exists():
        return ""
    items = []
    for path in sorted(processed_dir.rglob("*.txt")):
        try:
            stat = path.stat()
        except OSError:
            continue
        rel = path.relative_to(processed_dir).as_posix()
        items.append(f"{rel}:{stat.st_mtime_ns}:{stat.st_size}")
    if not items:
        return ""
    payload = "\n".join(items).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _build_index_manifest(
    processed_dir: Path,
    embedding_model: str,
    chunk_size: int,
    chunk_overlap: int,
) -> dict:
    return {
        "fingerprint": _fingerprint_processed(processed_dir),
        "embedding_model": embedding_model,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
    }


def _manifest_matches(stored: dict, expected: dict) -> bool:
    if not stored:
        return False
    for key, value in expected.items():
        if stored.get(key) != value:
            return False
    return True


def _save_manifest(persist_dir: Path, manifest: dict, doc_count: int) -> None:
    persist_dir.mkdir(parents=True, exist_ok=True)
    manifest_payload = dict(manifest)
    manifest_payload["doc_count"] = doc_count
    (persist_dir / "manifest.json").write_text(
        json.dumps(manifest_payload, ensure_ascii=False),
        encoding="utf-8",
    )


def _load_manifest(persist_dir: Path) -> dict:
    manifest_path = persist_dir / "manifest.json"
    if not manifest_path.exists():
        return {}
    try:
        return json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def _qdrant_location(persist_dir: Path) -> str:
    url = os.getenv("QDRANT_URL")
    if url:
        return url
    return os.getenv("QDRANT_PATH", str(persist_dir / "qdrant"))


def _get_qdrant_client(persist_dir: Path) -> QdrantClient:
    url = os.getenv("QDRANT_URL")
    api_key = os.getenv("QDRANT_API_KEY")
    if url:
        return QdrantClient(url=url, api_key=api_key)
    path = _qdrant_location(persist_dir)
    return QdrantClient(path=path)


def _collection_exists(client: QdrantClient, name: str) -> bool:
    try:
        client.get_collection(collection_name=name)
    except Exception:
        return False
    return True


def _recreate_collection(client: QdrantClient, name: str, vector_size: int) -> None:
    try:
        client.delete_collection(collection_name=name)
    except Exception:
        pass
    client.create_collection(
        collection_name=name,
        vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
    )


def _upsert_documents(
    client: QdrantClient,
    collection_name: str,
    docs: list[Document],
    vectors: list[list[float]],
    batch_size: int = 64,
) -> None:
    total = len(docs)
    for start in range(0, total, batch_size):
        points: list[PointStruct] = []
        for idx in range(start, min(start + batch_size, total)):
            payload = {
                "page_content": docs[idx].page_content,
                "metadata": docs[idx].metadata,
            }
            points.append(
                PointStruct(id=idx, vector=vectors[idx], payload=payload)
            )
        if points:
            client.upsert(collection_name=collection_name, points=points, wait=True)


def _search_qdrant(
    client: QdrantClient,
    collection_name: str,
    query_vector: list[float],
    limit: int,
):
    if hasattr(client, "search"):
        return client.search(
            collection_name=collection_name,
            query_vector=query_vector,
            limit=limit,
            with_payload=True,
        )
    if hasattr(client, "query_points"):
        response = client.query_points(
            collection_name=collection_name,
            query=query_vector,
            limit=limit,
            with_payload=True,
        )
        return getattr(response, "points", response)
    raise RuntimeError("Unsupported Qdrant client: missing search/query_points.")


def _docs_from_search_results(results) -> tuple[list[Document], np.ndarray]:
    docs: list[Document] = []
    scores: list[float] = []
    for point in results:
        payload = point.payload or {}
        content = payload.get("page_content") or ""
        metadata = payload.get("metadata") or {}
        if not isinstance(metadata, dict):
            metadata = {}
        docs.append(Document(page_content=content, metadata=metadata))
        scores.append(float(point.score or 0.0))
    return docs, np.array(scores, dtype=np.float32)


def _normalize_scores(scores: np.ndarray) -> np.ndarray:
    if scores.size == 0:
        return scores
    min_score = float(np.min(scores))
    max_score = float(np.max(scores))
    if max_score == min_score:
        return np.zeros_like(scores)
    return (scores - min_score) / (max_score - min_score)


def _keyword_scores_for_docs(
    docs: list[Document],
    cjk_keywords: list[str],
    latin_keywords: list[str],
) -> np.ndarray:
    scores = np.zeros(len(docs), dtype=np.float32)
    for idx, doc in enumerate(docs):
        scores[idx] = _keyword_score(doc.page_content, cjk_keywords, latin_keywords)
    return scores


def _hybrid_rerank(
    query: str,
    docs: list[Document],
    vector_scores: np.ndarray,
    k: int,
    alpha: float,
) -> list[Document]:
    if vector_scores.size == 0 or not docs:
        return []
    cjk, latin = _query_terms(query)
    keyword_scores = _keyword_scores_for_docs(docs, cjk, latin)

    if keyword_scores.size > 0 and float(keyword_scores.max()) > 0:
        combined = _normalize_scores(vector_scores) * alpha + _normalize_scores(
            keyword_scores
        ) * (1 - alpha)
        order = np.argsort(combined)[::-1][:k]
    else:
        order = np.argsort(vector_scores)[::-1][:k]

    return [docs[idx] for idx in order]


def _format_source(doc: Document) -> str:
    source = doc.metadata.get("source") or "unknown"
    name = Path(source).name
    page = doc.metadata.get("page")
    if page is not None:
        return f"{name}#p{page}"
    return name


def _format_context(docs: list[Document]) -> str:
    blocks = []
    for doc in docs:
        label = _format_source(doc)
        blocks.append(f"[{label}]\n{doc.page_content}")
    return "\n\n".join(blocks)


def _unique_sources(docs: list[Document]) -> list[str]:
    seen = set()
    sources: list[str] = []
    for doc in docs:
        label = _format_source(doc)
        if label not in seen:
            seen.add(label)
            sources.append(label)
    return sources


def _response_text(content) -> str:
    if content is None:
        return ""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict):
                text = item.get("text")
                if isinstance(text, str):
                    parts.append(text)
        return "\n".join(part for part in parts if part).strip()
    return str(content)


def _doc_signature(doc: Document) -> tuple:
    source = doc.metadata.get("source")
    page = doc.metadata.get("page")
    content_hash = hashlib.sha256(doc.page_content.encode("utf-8")).hexdigest()
    return (source, page, content_hash)


def _merge_docs(primary: list[Document], secondary: list[Document]) -> list[Document]:
    seen = set()
    merged: list[Document] = []
    for doc in primary + secondary:
        key = _doc_signature(doc)
        if key in seen:
            continue
        seen.add(key)
        merged.append(doc)
    return merged
